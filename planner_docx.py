from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path
from urllib.parse import quote
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT


TEMPLATE_PATH = Path(__file__).resolve().parent / "assets" / "templates" / "lesson-plan-template.docx"

DAY_KEYS = ["monday", "tuesday", "wednesday", "thursday", "friday"]
DAY_LABELS = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
DAY_COLUMNS = [0, 3, 5, 7, 9]
CENTER_LABELS = {
    "dramatic_play": "Dramatic Play",
    "construction": "Construction",
    "music": "Music",
    "art": "Art",
    "writing": "Writing",
    "manipulative_center": "Manipulative Center",
    "science": "Science",
    "sensory": "Sensory",
    "language_literacy": "Language & Literacy",
}
ATTACHMENT_ORDER = (
    [(f"circle_time.{key}", f"Circle Time - {label}") for key, label in zip(DAY_KEYS, DAY_LABELS)]
    + [(f"small_group_1.{key}", f"Small Group Learning Experience 1 - {label}") for key, label in zip(DAY_KEYS, DAY_LABELS)]
    + [(f"small_group_2.{key}", f"Small Group Learning Experience 2 - {label}") for key, label in zip(DAY_KEYS, DAY_LABELS)]
    + [("outdoor_learning", "Outdoor Learning Experiences")]
    + [(f"centers.{key}", f"Centers - {label}") for key, label in CENTER_LABELS.items()]
)


def _activity_item(plan: dict[str, object], section_name: str, key: str) -> dict[str, object]:
    activities = plan.get("activities") or {}
    if not isinstance(activities, dict):
        return {}
    section = activities.get(section_name) or {}
    if not isinstance(section, dict):
        return {}
    item = section.get(key) or {}
    return item if isinstance(item, dict) else {}


def _activity_text(plan: dict[str, object], section_name: str, key: str) -> str:
    return str(_activity_item(plan, section_name, key).get("text") or "")


def _activity_assessment(plan: dict[str, object], section_name: str, key: str) -> str:
    return str(_activity_item(plan, section_name, key).get("assessment") or "")


def _week_label(plan: dict[str, object]) -> str:
    month_label = str(plan.get("monthLabel") or "").strip()
    week_number = str(plan.get("weekNumber") or "").strip()
    if month_label and week_number:
        return f"{month_label} - Week {week_number}"
    if week_number:
        return f"Week {week_number}"
    return ""


def _safe_filename(filename: str) -> str:
    value = Path(filename or "attachment").name.strip() or "attachment"
    value = re.sub(r"[\r\n\t]", " ", value)
    return re.sub(r'[<>:"/\\|?*]', "_", value)


def _attachment_label(field_key: str) -> str:
    for key, label in ATTACHMENT_ORDER:
        if key == field_key:
            return label
    return field_key.replace(".", " - ").replace("_", " ").title()


def _decode_data_url(value: str) -> bytes:
    if not value.startswith("data:"):
        return b""
    _, _, payload = value.partition(",")
    if not payload:
        return b""
    return base64.b64decode(payload)


def _attachment_items(plan: dict[str, object]) -> list[dict[str, object]]:
    attachments = plan.get("attachments") or {}
    if not isinstance(attachments, dict):
        return []

    ordered_keys = [key for key, _ in ATTACHMENT_ORDER if key in attachments]
    ordered_keys.extend(sorted(key for key in attachments if key not in ordered_keys))

    result = []
    used_names: set[str] = set()
    for field_key in ordered_keys:
        item = attachments.get(field_key) or {}
        if not isinstance(item, dict):
            continue
        content = _decode_data_url(str(item.get("dataUrl") or ""))
        if not content:
            continue
        filename = _safe_filename(str(item.get("filename") or "attachment"))
        base_name = filename
        counter = 2
        while filename.lower() in used_names:
            stem = Path(base_name).stem or "attachment"
            suffix = Path(base_name).suffix
            filename = f"{stem}-{counter}{suffix}"
            counter += 1
        used_names.add(filename.lower())
        result.append(
            {
                "fieldKey": field_key,
                "label": _attachment_label(field_key),
                "filename": filename,
                "mimeType": str(item.get("mimeType") or "application/octet-stream"),
                "content": content,
            }
        )
    return result


def _set_cell_text(table: object, row: int, column: int, text: object) -> None:
    table.cell(row, column).text = str(text or "")


def _fill_day_section(table: object, plan: dict[str, object], section_name: str, text_row: int, assessment_row: int) -> None:
    for day_key, column in zip(DAY_KEYS, DAY_COLUMNS):
        _set_cell_text(table, text_row, column, _activity_text(plan, section_name, day_key))
        assessment = _activity_assessment(plan, section_name, day_key)
        _set_cell_text(table, assessment_row, column, f"Assessment: {assessment}" if assessment else "Assessment:")


def _fill_centers(table: object, plan: dict[str, object]) -> None:
    first_row = [
        ("dramatic_play", 15, 16, 0),
        ("construction", 15, 16, 3),
        ("music", 15, 16, 5),
        ("art", 15, 16, 7),
        ("writing", 15, 16, 9),
    ]
    second_row = [
        ("manipulative_center", 18, 19, 0),
        ("science", 18, 19, 3),
        ("sensory", 18, 19, 5),
        ("language_literacy", 18, 19, 7),
    ]
    for key, text_row, assessment_row, column in first_row + second_row:
        _set_cell_text(table, text_row, column, _activity_text(plan, "centers", key))
        assessment = _activity_assessment(plan, "centers", key)
        _set_cell_text(table, assessment_row, column, f"Assessment: {assessment}" if assessment else "Assessment:")


def _add_hyperlink(paragraph: object, url: str, text: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_attachments_section(document: object, items: list[dict[str, object]]) -> None:
    if not items:
        return
    document.add_page_break()
    document.add_heading("Attachments", level=1)
    document.add_paragraph("Attachment files are included in this Word file and listed below in lesson-plan order.")
    for item in items:
        paragraph = document.add_paragraph(style=None)
        paragraph.add_run(f"{item['label']}: ")
        data_url = (
            f"data:{quote(str(item['mimeType']), safe='/')};base64,"
            f"{base64.b64encode(bytes(item['content'])).decode('ascii')}"
        )
        _add_hyperlink(paragraph, data_url, str(item["filename"]))


def _add_embedded_attachment_files(docx_bytes: bytes, items: list[dict[str, object]]) -> bytes:
    if not items:
        return docx_bytes
    source = BytesIO(docx_bytes)
    target = BytesIO()
    with ZipFile(source, "r") as input_zip, ZipFile(target, "w", ZIP_DEFLATED) as output_zip:
        existing = set(input_zip.namelist())
        for name in input_zip.namelist():
            output_zip.writestr(name, input_zip.read(name))
        for index, item in enumerate(items, start=1):
            entry_name = f"word/attachments/{index:02d}-{item['filename']}"
            if entry_name not in existing:
                output_zip.writestr(entry_name, bytes(item["content"]))
    return target.getvalue()


def output_filename(plan: dict[str, object]) -> str:
    month_label = str(plan.get("monthLabel") or "Month").strip().replace("/", "-")
    week_number = str(plan.get("weekNumber") or "Week").strip()
    class_name = str(plan.get("className") or "Class").strip().replace("/", "-")
    return f"Lesson Plan - {month_label} - Week {week_number} - {class_name or 'Class'}.docx"


def build_planner_docx(plan: dict[str, object]) -> tuple[bytes, str]:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template DOCX was not found at {TEMPLATE_PATH}")

    document = Document(str(TEMPLATE_PATH))
    if not document.tables:
        raise ValueError("Template DOCX does not contain the lesson-plan table.")
    table = document.tables[0]

    _set_cell_text(table, 0, 2, _week_label(plan))
    _set_cell_text(table, 0, 6, plan.get("className") or "")
    _set_cell_text(table, 0, 10, plan.get("programName") or "")
    _set_cell_text(table, 1, 1, plan.get("books") or "")

    _fill_day_section(table, plan, "circle_time", 4, 5)
    _fill_day_section(table, plan, "small_group_1", 7, 8)
    _fill_day_section(table, plan, "small_group_2", 9, 10)

    outdoor_text = str(plan.get("outdoorLearning") or "")
    outdoor_assessment = str(plan.get("outdoorAssessment") or "")
    if outdoor_assessment:
        outdoor_text = f"{outdoor_text}\nAssessment: {outdoor_assessment}".strip()
    _set_cell_text(table, 12, 0, outdoor_text)

    _fill_centers(table, plan)

    attachment_items = _attachment_items(plan)
    _append_attachments_section(document, attachment_items)

    buffer = BytesIO()
    document.save(buffer)
    result = _add_embedded_attachment_files(buffer.getvalue(), attachment_items)
    return result, output_filename(plan)
